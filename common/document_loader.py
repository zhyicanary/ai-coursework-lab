"""多格式文档解析器 — 支持 PDF / DOCX / TXT / MD。

提供统一的 load() 入口，返回结构化的 Document 列表。
"""

import re
from pathlib import Path
from typing import IO


class Document:
    """单个文档片段，保留来源元信息。"""

    def __init__(self, content: str, metadata: dict | None = None):
        self.content = content
        self.metadata = metadata or {}

    def __repr__(self) -> str:
        return f"Document(content_len={len(self.content)}, meta={self.metadata})"


def load(file: str | Path | IO, filename: str | None = None) -> list[Document]:
    """加载文档文件，返回 Document 列表（每页/每段一个元素）。

    Args:
        file: 文件路径或文件对象（BytesIO / 上传文件）。
        filename: 文件名（当 file 是 IO 对象时用于判断格式）。

    Returns:
        list[Document]: 解析后的文档片段列表。
    """
    if isinstance(file, (str, Path)):
        path = Path(file)
        ext = path.suffix.lower()
        raw_text = _read_path(path)
        source = path.name
    else:
        ext = Path(filename or "unknown").suffix.lower()
        raw = file.read()
        if isinstance(raw, bytes):
            raw_text = raw.decode("utf-8", errors="replace")
        else:
            raw_text = raw
        source = filename or "upload"

    if ext == ".pdf":
        docs = _load_pdf(raw_text)
    elif ext in (".docx", ".doc"):
        docs = _load_docx(raw_text)
    elif ext in (".md", ".markdown"):
        docs = _load_markdown(raw_text)
    else:  # .txt 及其他纯文本
        docs = _load_text(raw_text)

    for d in docs:
        d.metadata.setdefault("source", source)
    return docs


def chunk_documents(
    documents: list[Document],
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> list[Document]:
    """将 Document 列表按字符粒度分块。

    Args:
        documents: 原始文档列表。
        chunk_size: 每块字符数。
        chunk_overlap: 块间重叠字符数。

    Returns:
        分块后的 Document 列表，每块保留所属文件的元信息。
    """
    chunks: list[Document] = []
    for doc in documents:
        text = doc.content
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            # 尽量在句子边界断开
            if end < len(text):
                # 往后找最近的句号/换行
                truncated = chunk_text.rfind("。")
                if truncated > chunk_size // 2:
                    chunk_text = chunk_text[: truncated + 1]
                    end = start + truncated + 1
                else:
                    nl = chunk_text.rfind("\n")
                    if nl > chunk_size // 2:
                        chunk_text = chunk_text[: nl + 1]
                        end = start + nl + 1

            meta = dict(doc.metadata)
            meta["chunk_index"] = len(chunks)
            chunks.append(Document(content=chunk_text.strip(), metadata=meta))
            start = end - chunk_overlap
            if start >= len(text):
                break
    return chunks


def list_supported_extensions() -> list[str]:
    return [".pdf", ".docx", ".doc", ".md", ".markdown", ".txt"]


# ── 内部读取 ──────────────────────────────────────────────


def _read_path(path: Path) -> str:
    """读取文件内容为文本。"""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _read_pdf_raw(path)
    elif ext in (".docx", ".doc"):
        return _read_docx_raw(path)
    else:
        return path.read_text("utf-8", errors="replace")


def _read_pdf_raw(path: Path) -> str:
    """用 PyPDF2 提取 PDF 文本。"""
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _read_docx_raw(path: Path) -> str:
    """用 python-docx 提取 DOCX 文本。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paras)


# ── 文本解析 ──────────────────────────────────────────────


def _load_pdf(raw: str) -> list[Document]:
    pages = re.split(r"\n\s*\n", raw)
    return [Document(content=p.strip()) for p in pages if p.strip()]


def _load_docx(raw: str) -> list[Document]:
    paras = [p.strip() for p in raw.split("\n\n") if p.strip()]
    return [Document(content=p) for p in paras]


def _load_markdown(raw: str) -> list[Document]:
    """按标题分割 Markdown 文档。"""
    sections = re.split(r"(?=^#{1,6}\s+)", raw, flags=re.MULTILINE)
    docs = []
    for sec in sections:
        sec = sec.strip()
        if sec:
            docs.append(Document(content=sec))
    if not docs:
        docs.append(Document(content=raw.strip()))
    return docs


def _load_text(raw: str) -> list[Document]:
    lines = [l.strip() for l in raw.split("\n") if l.strip()]
    return [Document(content=l) for l in lines] if lines else [Document(content=raw.strip())]
